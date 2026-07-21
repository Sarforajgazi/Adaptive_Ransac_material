import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
doc = Document('Adaptive_RANSAC_RL_Report.docx')

def find_para(prefix, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i >= start and p.text.strip() == prefix:
            return i
    return None

idx = find_para("5. Datasets")
if idx is not None and idx + 1 < len(doc.paragraphs):
    target_p = doc.paragraphs[idx + 1]
    
    intro_text = (
        "This project utilises two distinct datasets to train and evaluate the reinforcement learning policy. "
        "The TartanAir synthetic dataset serves as the primary source for RL training and baseline comparison, "
        "providing a vast array of geometric configurations across multiple simulated environments. "
        "The RELLIS-3D real-world dataset is employed strictly as a zero-shot generalisation test, "
        "providing human-annotated ground-truth labels for completely unseen physical terrain."
    )
    
    p_intro = target_p.insert_paragraph_before(intro_text)
    p_intro.style = 'Normal'
    print('intro added successfully.')
else:
    print('Failed to find section heading.')

doc.save('Adaptive_RANSAC_RL_Report.docx')
