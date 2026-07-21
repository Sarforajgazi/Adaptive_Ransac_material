import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml.ns import qn

doc = Document('Adaptive_RANSAC_RL_Report.docx')

idx_8 = None
for i, p in enumerate(doc.paragraphs):
    # Only match if it's the actual heading, not the TOC entry
    if p.text.strip().startswith("8. Discussion") and p.style and 'Heading 1' in p.style.name:
        idx_8 = i
        break

if idx_8 is not None:
    ref_element = doc.paragraphs[idx_8]._element
    body = ref_element.getparent()
    
    new_paragraphs = [
        ("Heading 2", "7.7 Zero-Shot Generalisation on Held-Out Environments"),
        ("Normal", (
            "To verify that the RL policy has learned generalisable geometric principles rather than "
            "memorising the training dataset, the v4 model was evaluated on six entire TartanAir environments "
            "(Gascola, House, WesternDesertTown, CoalMine, AbandonedFactory, NordicHarbor) that were "
            "strictly excluded from training."
        )),
        ("Normal", (
            "The RL agent successfully outperformed the strongest fixed baseline in every single unseen "
            "environment, confirming strong zero-shot generalisation capability even on completely novel topologies "
            "(such as the coastal water scenes of NordicHarbor)."
        )),
        ("Normal", (
            "Table 7.6 — Mean inlier ratio on held-out generalisation test environments (v4 model vs. baselines)."
        )),
        ("Normal", (
            "Gascola: RL (0.0631) > Standard (0.0474)  |  "
            "House: RL (0.1906) > Strict (0.1813)  |  "
            "WesternDesertTown: RL (0.1475) > Strict (0.1356)  |  "
            "CoalMine: RL (0.1367) > Standard (0.1194)  |  "
            "AbandonedFactory: RL (0.1318) > Strict (0.0884)  |  "
            "NordicHarbor: RL (0.0842) > Strict (0.0595)"
        ))
    ]
    
    for style_name, text in new_paragraphs:
        new_p = etree.Element(qn('w:p'))
        pPr = etree.SubElement(new_p, qn('w:pPr'))
        pStyle = etree.SubElement(pPr, qn('w:pStyle'))
        pStyle.set(qn('w:val'), style_name)
        run = etree.SubElement(new_p, qn('w:r'))
        t = etree.SubElement(run, qn('w:t'))
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        ref_element.addprevious(new_p)

doc.save('Adaptive_RANSAC_RL_Report.docx')

# Re-open to add images to the newly inserted elements cleanly
doc = Document('Adaptive_RANSAC_RL_Report.docx')
idx_8 = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("8. Discussion") and p.style and 'Heading 1' in p.style.name:
        idx_8 = i
        break

if idx_8 is not None:
    target_p = doc.paragraphs[idx_8]
    
    p_img1 = target_p.insert_paragraph_before("")
    p_img1.alignment = 1 # center
    run1 = p_img1.add_run()
    run1.add_picture('plots/heldout_inlier_ratio.png', width=Inches(6.0))
    
    p_cap1 = target_p.insert_paragraph_before("Figure 7.1: Mean Inlier Ratio on Held-Out Environments (Higher is better)")
    p_cap1.alignment = 1 # center
    p_cap1.style = 'Normal'
    
    p_img2 = target_p.insert_paragraph_before("")
    p_img2.alignment = 1 # center
    run2 = p_img2.add_run()
    run2.add_picture('plots/heldout_bad_frames.png', width=Inches(6.0))
    
    p_cap2 = target_p.insert_paragraph_before("Figure 7.2: Bad Frame Rate (<5% inliers) on Held-Out Environments (Lower is better)")
    p_cap2.alignment = 1 # center
    p_cap2.style = 'Normal'

doc.save('Adaptive_RANSAC_RL_Report.docx')
print('Perfectly inserted 7.7 and images before Section 8 heading.')
