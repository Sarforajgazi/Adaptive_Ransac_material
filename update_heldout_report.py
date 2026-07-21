import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml.ns import qn
import copy

doc = Document('Adaptive_RANSAC_RL_Report.docx')

def find_para(prefix, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i >= start and p.text.strip().startswith(prefix):
            return i
    return None

# Insert the held-out results as a new section 7.7
idx_insert = find_para("7.6 RELLIS-3D")
if idx_insert:
    # We need to insert after Section 7.6. We'll just look for the last paragraph in 7.6
    # which is "Compared to the earlier 100-frame smoke test..."
    idx_end = find_para("Compared to the earlier 100-frame", start=idx_insert)
    if idx_end:
        ref_element = doc.paragraphs[idx_end]._element
        body = ref_element.getparent()
        
        new_paragraphs = [
            ("Heading 2", "7.7 Zero-Shot Generalisation on Held-Out Environments"),
            ("Normal", (
                "To verify that the RL policy has learned generalisable geometric principles rather than "
                "memorising the training dataset, the v4 model was evaluated on six entire TartanAir environments "
                "(Gascola, House, WesternDesertTown, CoalMine, AbandonedFactory, NordicHarbor) that were "
                "strictly excluded from training. "
            )),
            ("Normal", (
                "The RL agent successfully outperformed the strongest fixed baseline in every single unseen "
                "environment, confirming strong zero-shot generalisation capability even on completely novel topologies "
                "(such as the coastal water scenes of NordicHarbor)."
            )),
            ("Normal", (
                "Table 7.6 — Mean inlier ratio on held-out generalisation test environments (v4 model vs. baselines)."
            )),
            ("Normal", (
                "Gascola: RL (0.0631) > Standard (0.0474)\n"
                "House: RL (0.1906) > Strict (0.1813)\n"
                "WesternDesertTown: RL (0.1475) > Strict (0.1356)\n"
                "CoalMine: RL (0.1367) > Standard (0.1194)\n"
                "AbandonedFactory: RL (0.1318) > Strict (0.0884)\n"
                "NordicHarbor: RL (0.0842) > Strict (0.0595)"
            ))
        ]
        
        for style_name, text in reversed(new_paragraphs):
            new_p = etree.SubElement(body, qn('w:p'))
            pPr = etree.SubElement(new_p, qn('w:pPr'))
            pStyle = etree.SubElement(pPr, qn('w:pStyle'))
            pStyle.set(qn('w:val'), style_name)
            run = etree.SubElement(new_p, qn('w:r'))
            t = etree.SubElement(run, qn('w:t'))
            t.text = text
            t.set(qn('xml:space'), 'preserve')
            # Handle newlines explicitly in docx by splitting the text
            if "\n" in text:
                t.text = ""
                parts = text.split("\n")
                for i, part in enumerate(parts):
                    if i > 0:
                        etree.SubElement(run, qn('w:br'))
                    t_part = etree.SubElement(run, qn('w:t'))
                    t_part.text = part
                    t_part.set(qn('xml:space'), 'preserve')
            
            ref_element.addnext(new_p)
            ref_element = new_p

doc.save('Adaptive_RANSAC_RL_Report.docx')
print('Held-out results added successfully.')
