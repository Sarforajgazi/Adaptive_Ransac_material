import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
doc = Document('Adaptive_RANSAC_RL_Report.docx')

def find_para_containing(text):
    for i, p in enumerate(doc.paragraphs):
        if text in p.text:
            return i
    return None

def insert_table_before(p, rows, cols, data):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # Move the newly created table (at the end of the doc) to be before paragraph `p`
    p._p.addprevious(table._tbl)
    
    # Populate data
    for r in range(rows):
        for c in range(cols):
            table.cell(r, c).text = data[r][c]
            
    # Remove the original paragraph
    p._p.getparent().remove(p._p)
    return table

# --- TABLE 7.4 ---
idx_74 = find_para_containing("RELLIS3D_00000 (2,847 frames): IoU=0.3803")
if idx_74 is not None:
    data_74 = [
        ["Sequence", "Total Frames", "Mean IoU"],
        ["RELLIS3D_00000", "2,847", "0.3803"],
        ["RELLIS3D_00001", "2,319", "0.3696"],
        ["RELLIS3D_00002", "4,147", "0.4543"],
        ["RELLIS3D_00003", "2,184", "0.6165"],
        ["RELLIS3D_00004", "2,059", "0.6058"]
    ]
    insert_table_before(doc.paragraphs[idx_74], 6, 3, data_74)
    print("Inserted Table 7.4")

# --- TABLE 7.5 ---
idx_75 = find_para_containing("Table 7.5 — RELLIS-3D overall results")
if idx_75 is not None:
    p = doc.paragraphs[idx_75]
    p.text = "Table 7.5 — RELLIS-3D overall results (13,556 frames)"
    next_p = doc.paragraphs[idx_75 + 1]
    
    data_75 = [
        ["Model Version", "Mean IoU", "Precision", "Recall", "F1 Score"],
        ["v3 Model (Pre-Reward Fix)", "0.4692", "0.7201", "0.6250", "0.5817"],
        ["v4 Model (Post-Reward Fix)", "0.4734", "0.7126", "0.6391", "0.5847"]
    ]
    
    table = doc.add_table(rows=3, cols=5)
    table.style = 'Table Grid'
    next_p._p.addprevious(table._tbl)
    
    for r in range(3):
        for c in range(5):
            table.cell(r, c).text = data_75[r][c]
            
    print("Inserted Table 7.5")

# --- TABLE 7.6 ---
idx_76 = find_para_containing("Gascola: RL (0.0631) > Standard (0.0474)")
if idx_76 is not None:
    data_76 = [
        ["Environment (Held-Out)", "RL Agent (v4)", "Best Fixed Baseline", "Winning Mode"],
        ["Gascola", "0.0631", "0.0474 (Standard)", "RL"],
        ["House", "0.1906", "0.1813 (Strict)", "RL"],
        ["WesternDesertTown", "0.1475", "0.1356 (Strict)", "RL"],
        ["CoalMine", "0.1367", "0.1194 (Standard)", "RL"],
        ["AbandonedFactory", "0.1318", "0.0884 (Strict)", "RL"],
        ["NordicHarbor", "0.0842", "0.0595 (Strict)", "RL"]
    ]
    insert_table_before(doc.paragraphs[idx_76], 7, 4, data_76)
    print("Inserted Table 7.6")

doc.save('Adaptive_RANSAC_RL_Report.docx')
print("Successfully saved docx with tables!")
