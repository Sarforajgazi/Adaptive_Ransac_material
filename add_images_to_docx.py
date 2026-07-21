import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches

doc = Document('Adaptive_RANSAC_RL_Report.docx')

def find_para(prefix, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i >= start and p.text.strip().startswith(prefix):
            return i
    return None

idx = find_para("Gascola: RL (")
if idx is not None:
    target_p = doc.paragraphs[idx]
    
    # We will insert the paragraphs AFTER this target_p
    # docx python library handles insert_paragraph_before, so we find the NEXT paragraph
    # and insert before that. If it's the last paragraph, we just add.
    
    if idx + 1 < len(doc.paragraphs):
        next_p = doc.paragraphs[idx + 1]
        
        p_img1 = next_p.insert_paragraph_before("")
        p_img1.alignment = 1 # center
        run1 = p_img1.add_run()
        run1.add_picture('plots/heldout_inlier_ratio.png', width=Inches(6.0))
        
        p_cap1 = next_p.insert_paragraph_before("Figure 7.1: Mean Inlier Ratio on Held-Out Environments (Higher is better)")
        p_cap1.alignment = 1 # center
        p_cap1.style = 'Normal'
        
        p_img2 = next_p.insert_paragraph_before("")
        p_img2.alignment = 1 # center
        run2 = p_img2.add_run()
        run2.add_picture('plots/heldout_bad_frames.png', width=Inches(6.0))
        
        p_cap2 = next_p.insert_paragraph_before("Figure 7.2: Bad Frame Rate (<5% inliers) on Held-Out Environments (Lower is better)")
        p_cap2.alignment = 1 # center
        p_cap2.style = 'Normal'
        print('Images inserted successfully.')
    else:
        print('Target paragraph was the last one, appending.')
        p_img1 = doc.add_paragraph("")
        p_img1.alignment = 1
        run1 = p_img1.add_run()
        run1.add_picture('plots/heldout_inlier_ratio.png', width=Inches(6.0))
        
        p_cap1 = doc.add_paragraph("Figure 7.1: Mean Inlier Ratio on Held-Out Environments (Higher is better)")
        p_cap1.alignment = 1
        
        p_img2 = doc.add_paragraph("")
        p_img2.alignment = 1
        run2 = p_img2.add_run()
        run2.add_picture('plots/heldout_bad_frames.png', width=Inches(6.0))
        
        p_cap2 = doc.add_paragraph("Figure 7.2: Bad Frame Rate (<5% inliers) on Held-Out Environments (Lower is better)")
        p_cap2.alignment = 1
        print('Images appended successfully.')
else:
    print('Could not find the target paragraph!')

doc.save('Adaptive_RANSAC_RL_Report.docx')
