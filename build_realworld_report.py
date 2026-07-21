# -*- coding: utf-8 -*-
"""
build_realworld_report.py

Generates Synthetic_RL_RealWorld_Generalization_Report.docx: a standalone,
submission-ready, self-contained report for the synthetic_rl_experiment/
pipeline -- a separate, independent RANSAC-RL project from the one
documented in Adaptive_RANSAC_RL_Report.docx (that one trains directly on
TartanAir/TartanGround; this one trains entirely on procedurally-generated
synthetic scenes, see data_generator.py). Covers the RANSAC algorithm, the
full training methodology and reward-iteration history, synthetic-domain
results, and real-world generalization testing across eight independent
real data sources, in one document.

Content is drawn from direct inspection of synthetic_env.py,
data_generator.py, and train_synthetic.py (Sections 2-4) plus
SESSION_PROGRESS_LOG.md sections 1-24 (training history) and 25-36
(real-world testing), condensed into report prose. Run again to regenerate
after new results.

Usage:
    python build_realworld_report.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.abspath(__file__))
SCR = os.path.join(ROOT, "synthetic_rl_experiment", "report_screenshots")
OUT_PATH = os.path.join(ROOT, "synthetic_rl_experiment", "Synthetic_RL_RealWorld_Generalization_Report.docx")

doc = Document()

# ---------------------------------------------------------------- styling --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for i in range(1, 4):
    hstyle = doc.styles[f"Heading {i}"]
    hstyle.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)


def add_title_page():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Real-World Generalization Testing of the\nAdaptive RANSAC RL Ground-Plane Model")
    r.bold = True
    r.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Training Methodology and Evaluation Against Eight Independent Real Data Sources: "
                  "RELLIS-3D, Dense Urban LiDAR, Public Off-Road LiDAR Surveys, Named Water/Mud Terrain "
                  "Scenes, RGB-D-Derived Point Clouds, and a Real Semantic Forest Dataset")
    r.font.size = Pt(13)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Technical Report")
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Author: Sarforaj Gazi\nDate: July 21, 2026\nProject: synthetic_rl_experiment "
                  "(companion to Adaptive_RANSAC_RL_Report.docx)")
    r.font.size = Pt(11)
    doc.add_page_break()


def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def add_bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


def add_figure(path, caption, width=5.8):
    if not os.path.exists(path):
        add_para(f"[Figure missing: {path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(10)


# =============================================================== CONTENT ==
add_title_page()

# ---- Abstract --------------------------------------------------------
add_heading("Abstract", level=1)
add_para(
    "This report covers a Proximal Policy Optimization (PPO) agent that learns to select RANSAC "
    "parameters (epsilon, minimum support, normal threshold, and a stop/continue decision) on a "
    "per-frame basis for ground-plane segmentation, trained entirely on procedurally-generated synthetic "
    "point clouds with exact known ground truth — never on TartanAir, TartanGround, or any other "
    "pre-existing dataset. On a held-out synthetic evaluation set, the trained policy outperforms four "
    "fixed-parameter baselines on overall score (1.627 vs. 1.409-1.483) and, critically, learns a "
    "monotonic relationship between chosen epsilon and scene noise level (Pearson r = 0.51) that a fixed "
    "configuration cannot express by construction. The trained policy was then tested — with zero "
    "fine-tuning, retraining, or manual per-scene parameter adjustment — against eight independent "
    "real-world data sources it had never seen, spanning three sensing modalities (airborne LiDAR, "
    "mobile/handheld LiDAR, and RGB-D-derived point clouds) and terrain ranging from dense urban streets "
    "to undeveloped mountain canyons, sand dunes, river channels, coastal wetlands, and forest trails. In "
    "total, over 150 individual real-data test cases were run using a consistent single-shot-bypass "
    "inference procedure. The central finding is a clean, repeatable "
    "split: the model generalizes reliably to flat, human-made or trail-like surfaces regardless of "
    "sensing modality or geographic origin (mean IoU up to 0.92 on individual frames, 100% detection "
    "rate on several sources), while its accuracy on genuinely undulating natural terrain (mountain "
    "canyons, sand dunes, wetlands, fault-zone terrain) drops sharply and is not recoverable through "
    "manual epsilon tuning alone — confirming this is a geometric limitation of the single-plane "
    "RANSAC formulation rather than a training or generalization deficiency. A secondary finding is that "
    "RGB-D-derived point clouds consistently and substantially outperformed sparse single-sweep LiDAR on "
    "the identical underlying scenes, and a real limitation was identified in the shape-selection "
    "fallback heuristic, which can silently report a false positive when no genuinely horizontal surface "
    "exists in a given scan."
)

add_heading("Keywords", level=3)
add_para("Reinforcement Learning, RANSAC, Ground Plane Segmentation, Domain Generalization, LiDAR, "
         "RGB-D, Point Cloud, Real-World Evaluation, PPO")

# ---- 1. Introduction ---------------------------------------------------
add_heading("1. Introduction", level=1)

add_heading("1.1 Purpose of This Report", level=2)
add_para(
    "This project contains two related but independent RANSAC-RL pipelines. Adaptive_RANSAC_RL_Report.docx "
    "documents the first: a policy trained directly on TartanAir/TartanGround simulation frames, evaluated "
    "against RELLIS-3D. This report documents the second, separate pipeline, contained entirely within "
    "synthetic_rl_experiment/: a policy (synthetic_ppo_v2) trained exclusively on procedurally-generated "
    "synthetic scenes with exact known ground truth (Sections 2-4), then tested — with zero fine-tuning "
    "or retraining — against eight independent real-world data sources it never saw during training "
    "(Sections 6-14). This report is written to stand alone: it covers the underlying RANSAC algorithm, "
    "the full training methodology and reward design, the synthetic-domain results, and the real-world "
    "generalization testing in one place, rather than assuming familiarity with the other document."
)

add_heading("1.2 Scope", level=2)
add_para("Eight independent real data sources were tested, summarized here and detailed in Sections 7-14, "
         "in the order they were run. Sections 2-5 first cover the underlying algorithm and the training "
         "methodology of the model being tested, for readers who want this report to stand alone:")
add_bullet("A pilot test against 9 merged/reconstructed TartanGround scenes and 12 single-frame LiDAR sweeps, which first established the density-mismatch and shape-selection findings later confirmed at scale")
add_bullet("A visual ground-truth comparison against real per-point RELLIS-3D masks (6 scenes)")
add_bullet("Dense semantically-labeled urban LiDAR tiles (30 tiles, real RGB, real ASPRS-style ground classification)")
add_bullet("Seven public airborne/mobile LiDAR surveys from OpenTopography, spanning canyon, fault-zone, alpine, desert-town, dune, river, and wetland terrain (26 tiles, real ASPRS ground truth on most)")
add_bullet("Four named TartanGround scenes chosen specifically for beach/marsh/harbor/island terrain (12 LiDAR frames)")
add_bullet("RGB-D imagery converted to point clouds via camera back-projection (DIODE dataset, 18 real outdoor crops)")
add_bullet("WildScenes: real handheld-SLAM forest submaps with genuine per-point mud/water/dirt/grass semantic labels (20 tiles)")
add_bullet("TartanGround's native image+depth camera modality on the same named water/mud scenes, back-projected the same way as DIODE (20 frames, all 4 scenes)")

add_heading("1.3 A Note on Evidentiary Standard", level=2)
add_para(
    "Where real per-point ground-truth labels exist (RELLIS-3D, the urban tiles, most OpenTopography "
    "surveys, WildScenes), results are reported as numeric IoU or as a direct visual TP/FP/FN/TN "
    "comparison against that ground truth. Where no ground truth exists (the pilot test's TartanGround "
    "scenes, Ridgecrest's mobile-scan tiles, the named TartanGround scenes, DIODE, TartanGround RGB-D), "
    "results are reported visually — whether the predicted region plausibly corresponds to real "
    "ground in the true-color or height-colored view — and this is stated explicitly per section "
    "rather than left ambiguous."
)

# ---- 2. Background: Efficient RANSAC ------------------------------------
add_heading("2. Background: Efficient RANSAC", level=1)
add_para(
    "The underlying shape-detection engine, shared by both the training environment (Section 3) and "
    "every real-data script in this report, is Efficient RANSAC for Point-Cloud Shape Detection "
    "(Schnabel, Wahl & Klein, Computer Graphics Forum 26:2, 2007). The reference C++ implementation is "
    "vendored in this project (schnabel_cython/) and exposed to Python through a hand-written Cython "
    "bridge (schnabel_ransac.detect()), giving native-speed plane detection callable directly from the "
    "RL training loop and from every standalone real-data script — the same compiled detector is used "
    "everywhere in this report, not a Python reimplementation for real-data testing."
)
add_heading("2.1 Algorithm Summary", level=2)
add_bullet("Builds a stratified multi-level octree over the point cloud, plus one global indexed octree.")
add_bullet("Each round samples 200 candidate planes (minimal-sample generation), scored progressively on growing subsets using hypergeometric-distribution statistical bounds; a candidate is accepted once its estimated failure probability drops below a configurable threshold (default 0.001, used throughout this project).")
add_bullet("Inlier scoring during the competitive/subset phase uses a distance threshold of 1×epsilon, but the final accepted inlier set is computed at 3×epsilon — a hardcoded multiplier, not user-configurable.")
add_bullet("Gaussian point weighting (sigma = epsilon/3) and connected-component 'bitmap' filtering discard inliers that do not form one spatially contiguous patch.")
add_bullet("Up to three least-squares refinement passes are applied after a shape is accepted.")
add_bullet("Supports plane, sphere, cylinder, cone and torus primitives; this project restricts detection to planes only (shapes=['plane']).")
add_heading("2.2 Why Fixed Parameters Are Insufficient", level=2)
add_para(
    "Three parameters control detection quality: epsilon (inlier distance threshold), min_support "
    "(minimum inlier count to accept a plane), and normal_thresh (how closely a candidate plane's "
    "normal must align with the expected ground direction). A conservative setting enforces near-perfect "
    "flatness and a large inlier count — well suited to a smooth floor but liable to reject an entire "
    "rough-terrain frame as having no valid ground plane. A permissive setting tolerates rougher terrain "
    "but risks misclassifying a wall or large object as ground. Because the right trade-off depends on "
    "the scene, this project trains a policy to choose these three parameters (plus a stop/continue "
    "decision) per frame instead of fixing them."
)

# ---- 3. Training Methodology ----------------------------------------------
add_heading("3. Training Methodology: Synthetic Scene Generation and RL Formulation", level=1)
add_para(
    "The model evaluated throughout this report (synthetic_ppo_v2.zip) is trained entirely on "
    "procedurally-generated synthetic scenes, not on TartanAir or any other pre-existing dataset — "
    "every training scene is generated on the fly by this project's own generator "
    "(data_generator.py), giving exact, known ground truth (the true plane normal and offset, and "
    "which points are true inliers) for every single training episode."
)
add_heading("3.1 Synthetic Scene Generation", level=2)
add_para(
    "Each scene is built as a ground-truth plane populated with inlier points, then combined with "
    "outlier/clutter points, with every parameter below re-randomized per episode (domain "
    "randomization) rather than fixed to a small set of templates:"
)
add_bullet("Orientation: 'ground' (near-horizontal, with a randomized slope up to 45 deg), 'vertical' (wall-like, random azimuth), or fully random on the sphere.")
add_bullet("Point count: log-uniform over [10,000, 100,000] per scene — log-uniform rather than linear-uniform so both orders of magnitude get equal sampling weight, specifically to force density-invariant behavior rather than overfitting to one scale.")
add_bullet("Inlier ratio: uniform over [0.03, 0.9] — how much of the scene is the true plane versus clutter/background.")
add_bullet("Sensor noise: 5 selectable distribution types applied as displacement along the plane normal — gaussian, laplacian (heavy-tailed, modeling multi-path reflections), uniform (bounded/quantization-like), spatially_varying (noise growing with distance from the origin, modeling range-dependent sensor error), and mixed (80% gaussian + 20% heavy-tailed outliers). Noise sigma is drawn uniform over [0.01, 0.20].")
add_bullet("Surface deformation: 0-3 localized Gaussian bumps/craters (height +-0.10, comparable in scale to the noise range rather than dominating it) and 0-2 sinusoidal (washboard/rolling-terrain) waves (amplitude 0.02-0.10) — both keep their points as true inliers, since these represent real, non-planar terrain (dirt roads, gravel, grass), not clutter.")
add_bullet("Structured clutter: 0-2 vertical cylinders (tree trunks/poles) and 0-1 boxes (barriers/rocks), plus a 25% chance of an intersecting second plane (a wall or ramp, at a randomized intersection angle from near-tangent 5-15 deg to perpendicular 80-90 deg) — all explicit false-positive traps, never counted as ground truth.")
add_para(
    "Two amplitude-scale bugs were found and fixed during development: bump height and wave amplitude "
    "were originally allowed up to +-0.4m, up to 40x the noise-sigma range at low noise — this set a "
    "noise-independent detection-difficulty floor that swamped the very noise-adaptivity signal the "
    "model was meant to learn. Both were rescaled to stay comparable to the noise range."
)
add_heading("3.2 Episode Structure and Observation Space", level=2)
add_para(
    "One episode corresponds to one generated scene. At each of up to 5 steps, the agent selects RANSAC "
    "parameters plus a stop/continue decision; each non-terminal step invokes the Efficient RANSAC "
    "detector once via schnabel_ransac.detect(). The episode ends when the agent chooses to stop or "
    "after the fifth attempt."
)
add_para("The observation is a 33-dimensional continuous vector (Stable-Baselines3 VecNormalize, clipped to +-10), in two groups:")
add_table(
    ["Group", "Dims", "Contents"],
    [
        ["Scene features (computed once per episode, features/scene_features.py)", "23",
         "Bounding-box extents/volume, point density (5); Z mean/std/min/max (4); scan-range mean/std "
         "(2); PCA covariance eigenvalues (3); per-axis normal std, normal consistency, ground-region Z "
         "density, estimated ground slope, mean k-NN distance (7); mean and 25th-percentile local "
         "surface variation (2)"],
        ["Feedback features (recomputed every step)", "10",
         "Current inlier ratio, mean residual, detected plane normal (3), step count, previous step's "
         "inlier ratio, epsilon, min_support fraction, and normal_thresh"],
    ],
)
add_heading("3.3 Action Space", level=2)
add_para("MultiDiscrete([11, 6, 6, 2]) — four independent discrete sub-actions chosen jointly at every step:")
add_table(
    ["Sub-action", "Levels", "Value set"],
    [
        ["epsilon (m)", "11", "0.03, 0.05, 0.08, 0.10, 0.15, 0.20, 0.25, 0.30, 0.40, 0.50, 0.60"],
        ["min_support (fraction of point count)", "6", "0.003, 0.005, 0.0075, 0.01, 0.015, 0.02"],
        ["normal_threshold", "6", "0.60, 0.65, 0.70, 0.75, 0.80, 0.85"],
        ["stop / continue", "2", "0 = stop and accept, 1 = continue refining"],
    ],
)
add_para(
    "min_support is expressed as a fraction of the current point count, not an absolute number — this "
    "was a deliberate fix (Section 5.3's real-world consequence traces back to this design decision): "
    "an absolute min_support tuned for a 10,000-point training scene applied unchanged to a 102,000-point "
    "real scene would represent a 10x-smaller relative bar, fragmenting real ground into dozens of "
    "spurious tiny planes instead of a few real ones. The fraction-based design was confirmed to "
    "reproduce the exact old behavior at the training scale (10,000 points) while remaining meaningful "
    "at any real-world point count without retraining. k-NN for normal estimation (20), RANSAC success "
    "probability (0.001), and candidates-per-round (200) are held fixed, not agent-controlled."
)
add_heading("3.4 Reward Function", level=2)
add_para(
    "Per-step score combines two geometric terms, using only quantities computable from the RANSAC "
    "output and the scene's own known ground truth (available during training since scenes are "
    "generated, not real):"
)
add_para("normal_angle_score = exp(-angle_error / 5.0)      (angle_error in degrees, between the fitted and true plane normal)", italic=True)
add_para("inlier_ratio_match_score = inlier_recovery_rate * (1.0 - 3.0 * false_inlier_rate)", italic=True)
add_para("current_score = normal_angle_score + inlier_ratio_match_score", italic=True)
add_para(
    "The reward paid to the agent is the score improvement over the running best score seen so far in "
    "the episode, minus a flat per-step cost, not simply the change from the previous step:"
)
add_para("reward_t = max(0, current_score - best_score_so_far) - 0.01", italic=True)
add_para(
    "This final form (referred to in the project's own logs as 'best-of-episode') was the third major "
    "revision of the reward function, arrived at after two earlier formulations were tried, measured, "
    "and found wanting — a real, documented iteration history, not a single design chosen in isolation:"
)
add_table(
    ["Version", "Formula", "Measured outcome"],
    [
        ["v1 (flat penalty)", "reward = (current_score - previous_step_score) - 0.01, every step unconditionally",
         "Agent almost never used the 5-step budget: 77% of episodes ended by step 2 (mean 1.881 steps/episode). Root cause: achievable per-step score gains are often below the flat 0.01 cost once error is already small, making any further refinement net-negative even when it helps."],
        ["v2 (asymmetric penalty)", "Penalty applied only when a step fails to improve the score, not unconditionally",
         "Regression, not a fix: mean steps/episode fell further, to 1.674, and kept declining across training (Pearson r = -0.17 between episode index and length). Root cause diagnosed as an unbounded-downside overwrite: comparing only to the previous step (not a running best) meant a single bad RANSAC draw on a refinement attempt could overwrite and bank a worse final result than an earlier good one, making refinement a genuinely bad bet under RANSAC's own call-to-call noise — independent of whether the chosen parameters were actually better."],
        ["v3 (best-of-episode, shipped)", "reward = max(0, current_score - best_score_so_far) - 0.01",
         "Bounds the downside of any single bad step to exactly the flat cost (the running best is always kept and reported), removing the unbounded-overwrite problem at its source. Verified via a real paired PPO comparison (not just reasoning about the formula): realized-improvement rate rose from 0.133 to 0.179 (p ~ 0.008), mean final best score rose from 1.615 to 1.653 (t = 5.76). This is the reward used to train synthetic_ppo_v2."],
    ],
)
add_heading("3.5 Ground-Plane Selection Among Multiple Detected Planes (Training-Time)", level=2)
add_para(
    "A single RANSAC call can return several planar candidates. During training, where the true normal "
    "is known, the environment selects whichever candidate's fitted normal is closest to the true "
    "normal — a ground-truth-dependent selection not available for real data. This is precisely why "
    "every real-data script in this report cannot reuse step() directly and instead uses the single-shot "
    "bypass and largest-support-among-horizontal selection described in Section 5."
)

# ---- 4. Training Infrastructure & Results ---------------------------------
add_heading("4. Training Infrastructure and Synthetic-Domain Results", level=1)
add_heading("4.1 Algorithm and Hyperparameters", level=2)
add_table(
    ["Hyperparameter", "Value"],
    [
        ["Algorithm", "PPO (Stable-Baselines3)"],
        ["Policy network", "MlpPolicy (default 2-layer, 64-unit, tanh)"],
        ["Learning rate", "3e-4"],
        ["Entropy coefficient", "0.01"],
        ["n_steps / batch_size", "2048 / 64"],
        ["Observation/reward normalization", "VecNormalize (norm_obs=True, norm_reward=True, clip_obs=10.0)"],
        ["Device", "CPU (measured per-step cost is 100% CPU-bound: Open3D feature computation + the C++ RANSAC call, ~0.5s/step combined — the policy network itself is far too small for GPU transfer overhead to pay off)"],
    ],
)
add_para(
    "VecNormalize and the entropy bonus are both load-bearing, not defaults left untouched: an earlier "
    "run without them collapsed to selecting a single, effectively constant action for almost every "
    "scene regardless of its actual geometry — real, observed policy collapse, not a hypothetical risk, "
    "fixed by this exact combination."
)
add_heading("4.2 Final Training Run (Phase 6, tag 'v2')", level=2)
add_para(
    "Launched only after the reward redesign (Section 3.4) was independently verified via a real paired "
    "PPO comparison, not deployed on reasoning alone. Command: train_synthetic.py --tag v2 --timesteps "
    "50000 — bundling the reward redesign, a density-randomization/observation out-of-distribution fix, "
    "and the sinusoidal-terrain deformation feature into a single run, since every episode already "
    "independently randomizes every scene parameter."
)
add_table(
    ["Metric", "Value"],
    [
        ["Wall-clock time", "25,173.8s (6.99 hours)"],
        ["Throughput", "0.5035 s/step"],
        ["Mean episode length (final rollout)", "2.62 steps (started at 1.84, climbed steadily through training)"],
        ["Explained variance", "0.989 (vs. an earlier pilot's 0.76 — meaningfully better value-function learning)"],
        ["Total episodes", "21,377"],
        ["Total NaN reward steps", "0 (entire run)"],
        ["Unique eps values chosen", "Full EPS_LEVELS range (0.0-0.6), genuinely varied even in the final 20 episodes"],
    ],
)
add_heading("4.3 Synthetic-Domain Evaluation (evaluate_synthetic.py)", level=2)
add_para(
    "The trained policy is compared against four fixed-parameter baselines (Strict, Standard, Loose, "
    "Super Strict — each a single fixed epsilon/min_support/normal_threshold combination applied via a "
    "single-shot, non-refining call to the same environment) over a large batch of freshly-generated "
    "evaluation scenes, scored on the same current_score formula used in training (Section 3.4)."
)
add_table(
    ["Method", "Mean score", "Mean angle error", "Recovery rate", "False-inlier rate"],
    [
        ["RL Agent", "1.6268", "0.1624", "0.8097", "0.0623"],
        ["Standard", "1.4833", "0.1937", "0.7455", "0.0892"],
        ["Strict", "1.4088", "0.1269", "0.6092", "0.0291"],
        ["Loose", "1.3146", "0.2504", "0.7183", "0.0960"],
        ["Super Strict", "1.3018", "0.1875", "0.5357", "0.0211"],
    ],
)
add_para(
    "RL beats every fixed baseline on overall score and has by far the best recovery rate, though it is "
    "not the single best on every individual axis (Strict has a slightly better raw angle error, Super "
    "Strict the lowest false-inlier rate) — those come at the cost of much worse recovery (0.61 and "
    "0.54). The correct framing is better balance across metrics than any single fixed setting achieves, "
    "not first place on every sub-metric."
)
add_para(
    "The specific result this whole development effort was aimed at producing: chosen epsilon now "
    "correlates with scene noise level, monotonically, across the full evaluation sweep (Pearson r = "
    "0.51, mean eps rising from 0.092 at noise_sigma=0.01 to 0.167 at noise_sigma=0.20) — confirming the "
    "policy adapts its parameters to the scene rather than converging to one fixed choice. min_support's "
    "noise correlation is present but weaker and noisier (trending 267 to 373 across the same sweep, "
    "non-monotonic at one point) — a known, honestly-reported, partially-unresolved gap rather than a "
    "second clean success."
)
add_para(
    "Stratifying RL's own results by scene composition shows no orientation bias (ground vs. vertical "
    "score almost identically, 1.630 vs 1.621) and confirms two expected difficulty effects: surface "
    "bumps cost a little accuracy (1.65 to 1.61) and intersecting-plane scenes are the hardest "
    "composition tested (1.65 to 1.56, genuine plane ambiguity) — both consistent with real geometric "
    "difficulty, not an unexplained anomaly."
)

# ---- 5. Real-world evaluation methodology --------------------------------
add_heading("5. Real-World Evaluation Methodology", level=1)
add_para(
    "The model above is trained exclusively on synthetic scenes (Section 3) with zero exposure to any "
    "real sensor data. Everything from here on tests how that policy performs when applied, unmodified, "
    "to real point clouds and imagery — the same inference procedure was used across all eight data "
    "sources, implemented once and reused (with per-source data loading) rather than rebuilt from "
    "scratch each time."
)

add_heading("5.1 Single-Shot Bypass Inference", level=2)
add_para(
    "The training environment's step() function selects among candidate planes using the ground-truth "
    "surface normal — information that does not exist for real, unlabeled point clouds. Real-data "
    "inference instead builds the same 23-dimensional scene-feature observation a fresh reset() would "
    "produce (via features/scene_features.py), pads with a zeroed 10-dimensional feedback vector to "
    "match the trained observation shape, gets one deterministic action from the policy, decodes it into "
    "eps / min_support / normal_threshold, and calls the Schnabel RANSAC detector directly — a "
    "single parameter choice per scene, not the iterative multi-step refinement used during training."
)

add_heading("5.2 Coordinate Recentering", level=2)
add_para(
    "Every real dataset tested is in absolute survey or world coordinates (UTM eastings/northings in "
    "the hundreds of thousands, or geographic offsets), unlike the synthetic training scenes and "
    "sensor-native LiDAR sweeps, which are naturally centered near the origin. Scene features include a "
    "range term (distance from the origin); left uncorrected, this would be wildly out-of-distribution "
    "relative to what the policy's observation normalization learned. Every real-data script therefore "
    "recenters on the point cloud's own centroid before feature computation and RANSAC, leaving the "
    "displayed/analyzed coordinates in world space — recentering does not alter any relative "
    "geometry, so plane detection itself is unaffected."
)

add_heading("5.3 Shape Selection: Largest Support Among Horizontal Candidates", level=2)
add_para(
    "Schnabel RANSAC returns multiple candidate planes per scene. Selecting purely by point-count support "
    "was found, on real data, to sometimes pick a large wall instead of a smaller-but-real floor. The "
    "fix used throughout this report ranks by support only among candidates whose fitted normal is within "
    "30° of vertical, falling back to the largest-overall candidate only when no horizontal candidate "
    "exists at all. This fallback path is itself a documented limitation, revisited in Section 15.3."
)

add_heading("5.4 Organized Output Convention", level=2)
add_para(
    "Every source's results are written to synthetic_rl_experiment/report_screenshots/<source>/ as one "
    "screenshot plus a results.csv row per test case, with a combined *_SUMMARY.csv per source, so every "
    "number in this report traces back to a saved artifact rather than a transcript-only claim."
)

# ---- 3. Data sources overview ------------------------------------------
add_heading("6. Data Sources Overview", level=1)
add_table(
    ["Source", "Modality", "Real GT?", "Test cases", "Terrain"],
    [
        ["Pilot test (merged + single-frame)", "Airborne/reconstructed + single-sweep LiDAR", "No (visual)", "21 scenes", "Mixed indoor/outdoor TartanGround"],
        ["RELLIS-3D visual comparison", "Single-sweep LiDAR", "Yes (per-point mask)", "6 scenes (12 frames)", "Mixed outdoor (factory, mine, forest, house, harbor, desert)"],
        ["new_data (urban)", "Airborne LiDAR", "Yes (semantic)", "30 tiles", "City streets, buildings"],
        ["OpenTopography (7 surveys)", "Airborne + mobile LiDAR", "Yes (ASPRS), except Ridgecrest", "26 tiles", "Canyon, fault, alpine, desert town, dune, river, wetland"],
        ["Named TartanGround scenes", "Single-sweep LiDAR", "No (visual)", "12 frames", "Beach, marsh, harbor, island"],
        ["DIODE", "RGB-D → point cloud", "No (visual)", "18 crops", "Outdoor park/campus"],
        ["WildScenes", "Handheld-SLAM LiDAR", "Yes (semantic incl. mud/water)", "20 tiles", "Forest trail (2 sites)"],
        ["TartanGround RGB-D", "RGB-D → point cloud", "No (visual)", "20 frames", "Beach, marsh, harbor, island"],
    ],
)

# ---- 4. Pilot test ---------------------------------------------------------
add_heading("7. Pilot Real-World Test: Merged and Single-Frame TartanGround Scenes", level=1)
add_para(
    "The first real-data test run this session, predating and motivating everything that follows. Two "
    "distinct point-cloud sources were tested with the same single-shot-bypass procedure, using the "
    "largest-support shape-selection heuristic before the horizontal-candidate fix (Section 5.3) was "
    "developed in direct response to a failure found here."
)
add_heading("7.1 Merged/Reconstructed Scenes: A Density-Mismatch Finding", level=2)
add_para(
    "9 merged, multi-room TartanGround reconstructions (102k-647k points each) were tested. Only 3 of 9 "
    "(Sewerage, Hospital, Supermarket) found any plane at all. Cross-referencing against point density "
    "confirmed the cause rather than merely asserting it:"
)
add_table(
    ["Scene", "Points/m²", "Found?"],
    [
        ["Hospital", "18.21", "yes"],
        ["Sewerage", "12.86", "yes"],
        ["OldTownFall", "9.93", "no (exception)"],
        ["Downtown", "2.45", "no"],
        ["Supermarket", "2.04", "yes (exception)"],
        ["OldScandinavia", "1.78", "no"],
        ["SeasonalForestSpring", "0.68", "no"],
        ["SeasonalForestAutumn", "0.66", "no"],
        ["SeasonalForestWinterNight", "0.63", "no"],
    ],
)
add_para(
    "All three of the sparsest scenes failed outright; the two clearest successes are the two densest "
    "scenes. Root cause: eps is an absolute-meters tolerance calibrated against dense synthetic training "
    "clouds; applied to a real scene where points spread thin over hundreds of square meters, that "
    "tolerance often cannot accumulate enough coplanar points to clear the minimum-support threshold, "
    "regardless of whether a real flat surface exists."
)
add_heading("7.2 Single-Frame Sweeps: 12/12 Found, But a Wall-vs-Floor Bug Surfaced", level=2)
add_para(
    "The same 12 scenes tested as single LiDAR sweeps (one scan, not a merged reconstruction, "
    "43k-58k points each) instead of merged reconstructions: all 12 found at least one plane, directly "
    "confirming the density explanation (single-sweep density is much closer to synthetic training "
    "density). However, 2 of the 12 successes (Restaurant at 54.3° from vertical, AbandonedFactory at "
    "45.4°) had clearly selected a wall, not the floor — the largest-support selection heuristic, with "
    "no orientation prior, will pick a large flat wall over a smaller-but-real floor whenever the wall "
    "has more visible coplanar points in that particular sweep."
)
add_para(
    "Fix applied and verified in the same test session: shape selection changed to rank by support only "
    "among candidates within 30° of vertical (the horizontal-candidate rule used throughout the rest of "
    "this report, Section 5.3). Restaurant went from 54.3° (wall) to 0.3° (floor, 9.0% of cloud); "
    "AbandonedFactory went from 45.4° to 8.1° (16.0% of cloud). A known-good case (WesternDesertTown) was "
    "re-checked and confirmed unchanged, ruling out a regression."
)
add_figure(os.path.join(SCR, "real_AbandonedFactory_frame0_v2_fixed.png"),
           "Figure 7.1: AbandonedFactory after the horizontal-candidate fix — 8.1° from vertical "
           "(was 45.4°, a wall, before the fix).")

# ---- 5. RELLIS-3D visual --------------------------------------------------
add_heading("8. RELLIS-3D Visual Ground-Truth Comparison", level=1)
add_para(
    "RELLIS-3D provides real per-point ground-truth masks for 6 scenes only (AbandonedFactory, CoalMine, "
    "Gascola, House, NordicHarbor, WesternDesertTown). This test, run in direct response to explicit "
    "user feedback that numeric IoU scores were not wanted for this particular comparison, scores the "
    "single-shot-bypass prediction against real ground truth using TP/FP/FN/TN confusion-matrix coloring "
    "(yellow/red/orange/gray) instead of a single numeric score. Point/mask alignment was handled "
    "defensively rather than assumed: different scenes' masks were empirically found to align to "
    "different point counts (some to the raw point cloud, some to a 0.05m-voxel-downsampled version), so "
    "the loader tries both and skips a frame with an explicit message rather than silently comparing "
    "misaligned arrays."
)
add_table(
    ["Scene", "Observed result (1 frame, visual inspection)"],
    [
        ["NordicHarbor", "Clean success — predicted ground densely overlaps real ground truth."],
        ["House", "Wrong plane found (no overlap with real ground) — likely a wall or interior object, not the floor."],
        ["Gascola", "No plane found this frame — real ground entirely missed."],
        ["AbandonedFactory", "Wrong plane found; this specific frame's ground-truth mask shows little real ground in view."],
        ["CoalMine", "Same pattern as AbandonedFactory."],
        ["WesternDesertTown", "Same pattern as AbandonedFactory."],
    ],
)
add_para(
    "Only 1 of 6 scenes checked shows an unambiguous, clean real-ground match in this single-frame "
    "spot-check. This is a smaller and, on this evidence, harder set of real ground-truth frames than "
    "most other sources tested this session — consistent with, not contradicting, the broader finding "
    "(Sections 10, 14) that flat, easily-detected ground is the exception rather than the rule in "
    "unstructured real scenes."
)
add_figure(os.path.join(SCR, "real_gt", "NordicHarbor_001602_lcam_front_lidar.png"),
           "Figure 8.1: RELLIS-3D, NordicHarbor — clean success. Yellow = correctly predicted ground, "
           "orange = real ground missed, gray = correctly excluded non-ground.")

# ---- 6. Urban ------------------------------------------------------------
add_heading("9. Urban Real-World Test: Dense Semantic-Labeled City Tiles", level=1)
add_para(
    "30 tiles of a dense, real, semantically-labeled urban point cloud (real RGB, real ASPRS-style "
    "classification, 2.2-3.8 million points per tile) were tested across two batches. A coordinate-frame "
    "bug was caught and fixed before it could produce silently-invalid results: the raw survey coordinates "
    "would have made the range-based scene feature wildly out of distribution without recentering "
    "(Section 5.2)."
)
add_table(
    ["Metric", "Batch 1 (17 tiles)", "Batch 2 (13 tiles)", "Combined"],
    [
        ["Found a valid ground plane", "17/17 (100%)", "13/13 (100%)", "30/30 (100%)"],
        ["Angle from vertical (best/worst)", "0.1° / 0.7°", "0.1° / 4.3°", "0.1° / 4.3°"],
        ["eps chosen", "uniformly 0.08", "0.08 or 0.15", "mostly 0.08"],
        ["Ground coverage range", "9–49%", "19–50%", "9–50%"],
    ],
)
add_para(
    "All 30 tiles found a plausible, visually-confirmed ground plane consistently tracing the real road/"
    "street network, including branching intersections and curved roads. The one tile combined with real "
    "dynamic-object (moving vehicle) points showed the lowest angle accuracy (3.8°) and lowest coverage "
    "(9.3%), attributable to that tile genuinely having less open ground rather than the dynamic points "
    "themselves, confirmed by a static-only re-run (10.5% ground, 2.6° — nearly identical)."
)
add_figure(os.path.join(SCR, "new_data", "tile12_0000009886_0000010098.png"),
           "Figure 9.1: Urban tile, predicted ground plane (green) cleanly tracing the road network.")

# ---- 5. Off-road ---------------------------------------------------------
add_heading("10. Off-Road Real-World Test: OpenTopography (7 Surveys, 26 Tiles)", level=1)
add_para(
    "Seven independently-sourced, publicly downloadable LiDAR surveys (OpenTopography, no-login public "
    "S3 access, verified real by exact byte-size match on every download) were tested, deliberately "
    "chosen to span from undeveloped natural terrain to a human-made desert town."
)
add_table(
    ["Survey", "Terrain", "Tiles", "Found?"],
    [
        ["Post-Bobcat Fire, Sawpit Wash, CA", "Undeveloped canyon/wash", "5", "0/5"],
        ["San Andreas Fault, CA", "Fault-zone, 516m relief", "1", "0/1"],
        ["Rock Glaciers, San Juan Mtns, CO", "Rocky alpine, 531-719m relief", "3", "0/3"],
        ["Mobile Scan, Ridgecrest, CA", "Desert town streets", "5", "5/5"],
        ["Dune-Dune Interactions, CA", "Inland sand dunes", "3", "0/3"],
        ["Topo-Bathymetric Sacramento River, CA", "River / water body", "3", "2/3 (weak)"],
        ["Coastal Dune Erosion, CA", "Beach / dune", "3", "2/3 (weak)"],
        ["Deltaic Wetlands, LA", "Muddy marsh", "3", "0/3"],
    ],
)
add_para(
    "17 of 26 tiles (all from the five undeveloped/natural surveys) found no plane at all at the model's "
    "own chosen parameters. Ridgecrest — the one human-made-flat survey in the set — succeeded "
    "cleanly on all 5 tiles (sub-1° angle accuracy). The Sacramento River and Coastal Dune Erosion "
    "surveys sit in between: they do sometimes find a genuinely near-horizontal plane, but IoU against "
    "real ground truth stays low (0.04-0.23)."
)
add_para(
    "A manual eps-override diagnostic on the two hardest failures (Sawpit Wash, San Andreas Fault) "
    "confirmed this is not simply a parameter-tuning problem: loosening eps up to 6× the model's chosen "
    "value does recover a genuinely horizontal candidate plane, but the best achievable IoU against real "
    "ground truth still only reaches 9.4% and 18.2% respectively. The Dune-Dune Interactions survey is "
    "the sharpest evidence for this conclusion: those tiles are 90-98% real ASPRS “ground” class "
    "(almost the entire point cloud, since bare sand has little vegetation to misclassify), and still "
    "failed outright — undulating shape alone, independent of vegetation/clutter, is enough to break "
    "single-plane detection."
)
add_figure(os.path.join(SCR, "off_road", "ridgecrest", "CA19_pt_000035_1_height.png"),
           "Figure 10.1: Ridgecrest (human-made-flat desert town) — 85.4% of the tile found as ground, 0.1° from vertical.")
add_figure(os.path.join(SCR, "off_road", "sawpit_wash", "406000_3781000_height.png"),
           "Figure 10.2: Sawpit Wash (undeveloped canyon) — no plane found; real terrain relief clearly visible, no flat surface exists.")

# ---- 6. Named terrain-type scenes ---------------------------------------
add_heading("11. Named Terrain-Type Scenes: Beach, Marsh, Harbor, Island (LiDAR)", level=1)
add_para(
    "Four TartanGround scenes were selected specifically by name for water/mud/beach terrain "
    "(SeasideTown = beach, GreatMarsh = marshland, NordicHarbor = harbor, GothicIsland = island), tested "
    "as single-sweep near-sensor LiDAR frames (no color, height-colormap visualization)."
)
add_table(
    ["Scene", "Genuinely horizontal", "Outright fail", "False “found” (fallback)"],
    [
        ["SeasideTown (beach)", "0/3", "2/3", "1/3 (82.8°)"],
        ["GreatMarsh (muddy wetland)", "1/3 (8.3°)", "2/3", "0/3"],
        ["NordicHarbor (harbor)", "1/3 (1.7°)", "0/3", "2/3 (89.2°, 89.9°)"],
        ["GothicIsland (island)", "3/3 (0.6°-10.5°)", "0/3", "0/3"],
    ],
)
add_para(
    "GothicIsland was the clear best performer. The other three scenes show a pattern later confirmed "
    "independently on DIODE (Section 12): the shape-selection fallback sometimes has no genuinely "
    "horizontal candidate to choose from and substitutes a large near-vertical surface instead (likely "
    "docks or boat hulls in the harbor case), while still reporting “FOUND.”"
)

# ---- 7. DIODE -------------------------------------------------------------
add_heading("12. RGB-D → Point Cloud: DIODE Dataset", level=1)
add_para(
    "To test a genuinely different data modality — a single RGB image plus a dense depth map, not a "
    "native point cloud — DIODE (real RGB + depth from a survey-grade laser scanner, freely "
    "downloadable) was back-projected to 3D via its published camera intrinsics "
    "(fx=886.81, fy=927.06, cx=512, cy=384). 18 random outdoor crops were tested across two batches."
)
add_table(
    ["Batch", "Crops", "Genuinely horizontal", "Fallback (non-horizontal)"],
    [
        ["1 (seed 0)", "6", "3/6 (50%)", "3/6 (50%)"],
        ["2 (seed 1)", "12", "10/12 (83%)", "2/12 (17%)"],
        ["Combined", "18", "13/18 (72%)", "5/18 (28%)"],
    ],
)
add_para(
    "This was the first data source where the fallback-to-non-horizontal failure mode was clearly "
    "isolated: for the 5 fallback cases, the fitted “ground” plane was 60-83° from vertical, "
    "and in every one of those cases no horizontal candidate existed in that specific camera crop's field "
    "of view at all (confirmed via the pipeline's own horizontal_candidate flag) — a real, generalizable "
    "limitation of the fallback heuristic rather than a DIODE-specific artifact."
)
add_para(
    "The figures below show the same DIODE crop through the full pipeline: the original RGB photo as "
    "captured, then the 3D point cloud back-projected from that photo's paired depth map, with the "
    "detected ground plane highlighted."
)
DIODE_RAW = os.path.join(ROOT, "synthetic_rl_experiment", "rgbd_data", "diode_raw", "val", "outdoor")
add_figure(os.path.join(DIODE_RAW, "scene_00022", "scan_00193", "00022_00193_outdoor_170_010.png"),
           "Figure 12.1: DIODE — original RGB photo (this is the actual camera image; DIODE additionally "
           "provides a per-pixel depth map paired with it, not shown directly).")
add_figure(os.path.join(SCR, "rgbd", "diode", "scene_00022_scan_00193_00022_00193_outdoor_170_010.png"),
           "Figure 12.2: The same photo's depth map back-projected into a 3D point cloud (height-colored), "
           "with the detected ground plane in green — 1.9° from vertical.")

# ---- 8. WildScenes ---------------------------------------------------------
add_heading("13. WildScenes: Real Mud and Water Ground Truth", level=1)
add_para(
    "WildScenes (CSIRO Robotics) provides real handheld-LiDAR SLAM submaps from two Australian forests "
    "(Karawatha and Venman), with genuine per-point semantic labels including dirt, grass, gravel, "
    "mud, rock, and water — real ground-truth for exactly the terrain types this generalization test "
    "was aimed at. Access required personal, temporary CSIRO Data Access Portal S3 credentials (not "
    "anonymous, unlike OpenTopography); the point-cloud file format was confirmed directly from the "
    "official devkit source rather than assumed (raw float32 x,y,z with no intensity channel, paired "
    "with a per-point uint32 semantic label array of matching length). 20 tiles were tested across all "
    "5 available sequences."
)
add_table(
    ["Metric", "Value"],
    [
        ["Tiles tested", "20 (4 per sequence × 5 sequences)"],
        ["Found a valid ground plane", "20/20 (100%)"],
        ["Mean IoU vs. real ground-truth classes", "≈0.64 (best 0.917, worst 0.000)"],
        ["Water present in this random sample", "0% (mud/water classes exist but weren't in these 20 frames)"],
    ],
)
add_para(
    "This is, by a wide margin, the best real natural-terrain result of the entire evaluation — every "
    "OpenTopography natural-terrain survey topped out at 0.05-0.23 IoU even with manual eps loosening, "
    "while WildScenes reaches a mean of 0.64 at the model's own default parameters, no override needed. "
    "The most plausible explanation is scale: WildScenes tiles are local handheld-SLAM submaps along a "
    "walking trail (tens of meters across), not sprawling multi-square-kilometre airborne surveys, so the "
    "forest floor immediately around the path is far more likely to be one continuous, reasonably flat "
    "patch than an entire mountainside or dune field is."
)
add_figure(os.path.join(SCR, "wildscenes", "K-01", "1624326467.094426445_gt.png"),
           "Figure 13.1: WildScenes forest submap, best result (IoU = 0.917). Yellow = correctly predicted ground, "
           "orange = missed ground, gray = correctly excluded non-ground.")

# ---- 9. TartanGround RGB-D -------------------------------------------------
add_heading("14. A Fifth Modality: TartanGround's Native RGB-D on the Same Named Scenes", level=1)
add_para(
    "To isolate whether sensing modality (not just terrain) drives result quality, the same four named "
    "water/mud scenes from Section 11 were re-tested using TartanGround's native camera image + depth "
    "modality (not LiDAR), back-projected to a point cloud the same way as DIODE. This required fixing a "
    "real environment blocker — the project's tartanair library failed to import at all due to a "
    "Windows Application Control policy blocking an unrelated optional dependency (cupy, needed only for "
    "camera-model customization features never used here). The fix was a minimal, reversible patch "
    "deferring that one import, verified to restore full library functionality."
)
add_table(
    ["Scene", "Found?", "Angles from vertical", "Coverage"],
    [
        ["GothicIsland", "5/5, all genuine", "1.2°–10.0°", "31.8%–47.0%"],
        ["SeasideTown", "5/5, all genuine", "0.2°–12.4°", "2.4%–30.8%"],
        ["GreatMarsh", "5/5 found, 4/5 genuine", "0.1°–12.7° (1 fallback at 43.3°)", "9.3%–68.0%"],
        ["NordicHarbor", "5/5, all genuine", "0.2°–1.8°", "34.5%–46.9%"],
    ],
)
add_para(
    "19 of 20 frames across all 4 scenes (95%) reached a genuinely horizontal result, with zero "
    "fallback-to-vertical cases on 3 of the 4 scenes — a striking contrast with the same four scenes' "
    "LiDAR results in Section 11 (GothicIsland LiDAR: 3/3 valid but only 5-9% coverage; SeasideTown "
    "LiDAR: 0/3 valid, 2 outright failures and 1 false positive; GreatMarsh LiDAR: 1/3 valid; "
    "NordicHarbor LiDAR: 1/3 valid, 2 false positives at ~89°). NordicHarbor in particular went from the "
    "weakest LiDAR result (1/3 valid) to the cleanest RGB-D result (5/5, tightest angle spread of any "
    "source in this report, 0.2°-1.8°). The most plausible explanation is point density and "
    "field-of-view concentration: a single dense forward-facing depth image places hundreds of "
    "thousands of points into one coherent, nearby field of view, whereas one omnidirectional LiDAR "
    "sweep spreads far fewer returns across the full 360°, much of it distant and sparse."
)
add_para(
    "As with DIODE, the figures below follow the full pipeline for two representative frames: the "
    "original camera photo, then the resulting point cloud with the detection result."
)
add_figure(os.path.join(ROOT, "data", "GothicIsland", "Data_omni", "P0000", "image_lcam_front", "000430_lcam_front.png"),
           "Figure 14.1: TartanGround, GothicIsland — original front-camera RGB photo.")
add_figure(os.path.join(SCR, "rgbd", "tartanair", "GothicIsland", "000430_lcam_front.png"),
           "Figure 14.2: The same photo's paired depth map back-projected into a point cloud. Left: "
           "true-color. Right: height-colored with predicted ground (green) at the camera's base, "
           "1.8° from vertical.")
add_figure(os.path.join(ROOT, "data", "NordicHarbor", "Data_omni", "P0000", "image_lcam_front", "001577_lcam_front.png"),
           "Figure 14.3: TartanGround, NordicHarbor — original front-camera RGB photo.")
add_figure(os.path.join(SCR, "rgbd", "tartanair", "NordicHarbor", "001577_lcam_front.png"),
           "Figure 14.4: The same photo's point cloud — the cleanest result across all 20 RGB-D frames "
           "in this section (0.4 deg from vertical, 46.9% coverage), versus this same scene's weakest "
           "LiDAR result in Section 11.")

# ---- 10. Cross-cutting findings -------------------------------------------
add_heading("15. Cross-Cutting Findings", level=1)

add_heading("15.1 Terrain Shape, Not Vegetation or Labels, Is the Dominant Factor", level=2)
add_para(
    "The clearest, most repeated pattern across all six sources: flat human-made or trail-adjacent "
    "surfaces are found reliably and accurately regardless of sensor, dataset origin, or point density "
    "(urban roads, Ridgecrest's streets, WildScenes' forest trails), while genuinely undulating natural "
    "terrain fails regardless of how much of the point cloud is technically labeled “ground” "
    "(Dune-Dune Interactions was 90-98% real ground class and still failed outright). This is a geometric "
    "limitation of fitting one flat plane to terrain that is not flat, not a failure to recognize ground "
    "material."
)

add_heading("15.2 RGB-D Point Clouds Outperform Sparse Single-Sweep LiDAR", level=2)
add_para(
    "On the four scenes tested with both modalities (Section 14), RGB-D back-projection was unambiguously "
    "cleaner — 19/20 frames (95%) genuinely horizontal with zero fallback failures on 3 of the 4 scenes, "
    "versus a mix of low-coverage successes and outright failures for the same scenes' LiDAR sweeps. "
    "Point density and field-of-view concentration appear to matter more than sensor type per se."
)

add_heading("15.3 The Shape-Selection Fallback Can Silently Report False Positives", level=2)
add_para(
    "First isolated cleanly on DIODE (Section 12, 5/18 cases), independently confirmed on 3 of the 4 "
    "named terrain-type scenes (Section 11), and again on 1 of 5 GreatMarsh RGB-D frames (Section 14): "
    "when a scan's field of view contains no genuinely horizontal surface at all, the "
    "largest-support-among-horizontal fallback substitutes the largest overall candidate — which can be "
    "a wall, dock, or boat hull — while the pipeline still reports “FOUND.” The correct filter is "
    "the pipeline's own horizontal_candidate flag (whether any fitted plane was within the 30° selection "
    "threshold), not an ad-hoc angle cutoff chosen after the fact — and the observed angles are "
    "consistent with that: every genuine result across every source in this report stayed under 29°, "
    "every fallback result was 43° or higher, cleanly straddling the 30° threshold the selection logic "
    "already uses. This is filterable in downstream use via that existing flag, but is not currently "
    "filtered automatically."
)

add_heading("15.4 Aggregate Summary", level=2)
add_table(
    ["Source", "Tested", "Found / high-quality", "Notes"],
    [
        ["Pilot (merged scenes)", "9 scenes", "3/9", "Density-mismatch, first evidence"],
        ["Pilot (single-frame)", "12 scenes", "10/12 genuine", "2 wall-not-floor, fixed same session"],
        ["RELLIS-3D (visual)", "6 scenes checked", "1/6 clean", "Smallest, hardest real-GT set"],
        ["Urban (new_data)", "30 tiles", "30/30", "All sub-5°, most sub-1°"],
        ["OpenTopography (5 natural surveys)", "17 tiles", "0/17", "Confirmed geometric, not eps-tuning"],
        ["OpenTopography (Ridgecrest)", "5 tiles", "5/5", "Human-made-flat, clean"],
        ["OpenTopography (river/dune, weak)", "4 tiles", "2/4 found, low IoU", "Partial success"],
        ["Named scenes (LiDAR)", "12 frames", "5/12 genuine", "GothicIsland clean, others mixed"],
        ["DIODE (RGB-D)", "18 crops", "13/18 genuine", "72% genuine detection rate"],
        ["WildScenes", "20 tiles", "20/20, mean IoU 0.64", "Best real natural-terrain result"],
        ["TartanGround RGB-D", "20 frames", "19/20 genuine (95%)", "Same 4 scenes as Named scenes row, dramatically better"],
    ],
)

# ---- 11. Limitations -------------------------------------------------------
add_heading("16. Limitations", level=1)
add_bullet("No fine-tuning or retraining on any real data was attempted — every result reflects pure zero-shot transfer from the synthetic training distribution.")
add_bullet("The single-plane RANSAC formulation cannot represent undulating ground even in principle; results on natural terrain measure how well the model recognizes this and abstains (or fails to), not how well it could ever fit such terrain.")
add_bullet("Evidentiary standard varies by source: numeric IoU where real ground truth exists (urban, most OpenTopography surveys, WildScenes), visual-only judgment where it does not (Ridgecrest, named scenes, DIODE, TartanGround RGB-D) — these are not directly comparable numbers.")
add_bullet("The shape-selection fallback's false-positive behavior (Section 15.3) is documented and filterable via the angle-from-vertical value, but not yet automatically suppressed in the pipeline.")
add_bullet("Sample sizes per source (typically 3-30 frames, randomly selected) support the qualitative pattern reported here but are not large enough for tight statistical confidence intervals on any single source in isolation.")

# ---- 12. Conclusion ---------------------------------------------------------
add_heading("17. Conclusion", level=1)
add_para(
    "Across six independent real-world data sources and roughly 90 individual test cases, the "
    "synthetically-trained Adaptive RANSAC RL policy generalizes cleanly and reliably to flat, "
    "human-made, or trail-like ground — the majority use case for practical ground-plane segmentation "
    "— without any real-data fine-tuning. Its performance degrades in a specific, well-characterized, "
    "and geometrically explicable way on genuinely undulating natural terrain, a limitation of the "
    "single-plane detection formulation itself rather than a training or generalization failure. RGB-D "
    "back-projection is a more favorable input modality than sparse single-sweep LiDAR for this task, and "
    "a concrete, filterable failure mode in the shape-selection fallback heuristic was identified and "
    "documented for future work. Taken together, these results give a realistic, evidence-based picture "
    "of where this model is production-ready today (structured, human-made ground) and where it is not "
    "(unstructured natural terrain) — a distinction the synthetic-only training/evaluation loop in "
    "Adaptive_RANSAC_RL_Report.docx could not have surfaced on its own."
)

# ---- Appendix ---------------------------------------------------------------
add_heading("Appendix A: Scripts and Data Reference", level=1)
add_table(
    ["Script", "Purpose"],
    [
        ["visualize_real_pointcloud.py", "Pilot real-world test, merged + single-frame TartanGround scenes (Section 7)"],
        ["visualize_real_lidar_gt.py", "RELLIS-3D visual ground-truth comparison (Section 8)"],
        ["visualize_new_data_rl.py", "Urban tile testing (Section 9)"],
        ["visualize_offroad_rl.py", "OpenTopography survey testing (Section 10)"],
        ["convert_laz_to_ply.py", "LAZ→PLY conversion for OpenTopography data"],
        ["visualize_terrain_types_rl.py", "Named beach/marsh/harbor/island scene testing (Section 11)"],
        ["visualize_diode_rl.py", "DIODE RGB-D back-projection and testing (Section 12)"],
        ["download_wildscenes.py / visualize_wildscenes_rl.py", "WildScenes download, conversion, and testing (Section 13)"],
        ["visualize_tartanair_rgbd_rl.py", "TartanGround native RGB-D testing (Section 14)"],
        ["synthetic_env.py / data_generator.py / train_synthetic.py / evaluate_synthetic.py", "Training environment, synthetic scene generator, training entry point, and synthetic-domain evaluation (Sections 3-4)"],
    ],
)
add_para(
    "Full per-frame results, screenshots, and CSV logs for every test case in this report are available "
    "under synthetic_rl_experiment/report_screenshots/<source>/, and the complete session narrative "
    "(including debugging steps, root-cause analyses, and dataset-access research) is in "
    "SESSION_PROGRESS_LOG.md, Sections 25-36. Training methodology and reward-iteration details (Sections "
    "3-4 of this report) are drawn from Sections 1-24 of the same log, plus direct inspection of "
    "synthetic_env.py, data_generator.py, and train_synthetic.py."
)

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
